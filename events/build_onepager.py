#!/usr/bin/env python3
"""Build a CEO-ready one-page briefing for the DakCUPAC bourbon tasting fundraiser."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x1F, 0x2A, 0x44)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x3A, 0x3A, 0x3A)
LIGHT = RGBColor(0x88, 0x88, 0x88)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "1F4E79")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # half-points, so 9pt
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def tight(paragraph, space_before=0, space_after=2):
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.1


def add_heading(doc, text, size=11, color=BLUE, bold=True, space_before=6, space_after=2):
    p = doc.add_paragraph()
    tight(p, space_before=space_before, space_after=space_after)
    r = p.add_run(text.upper())
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_body(doc, text, size=9.5, bold=False, color=GRAY, space_before=0, space_after=2):
    p = doc.add_paragraph()
    tight(p, space_before=space_before, space_after=space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.font.name = "Calibri"
    return p


def add_bullet(doc, prefix_runs, size=9):
    """prefix_runs = list of (text, bold, is_link, url). A bullet glyph is prepended."""
    p = doc.add_paragraph()
    tight(p, space_before=0, space_after=1)
    p.paragraph_format.left_indent = Inches(0.12)
    bullet_run = p.add_run("•  ")
    bullet_run.font.size = Pt(size)
    bullet_run.font.color.rgb = BLUE
    bullet_run.bold = True
    for text, bold, is_link, url in prefix_runs:
        if is_link:
            add_hyperlink(p, url, text)
        else:
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.color.rgb = GRAY
            r.bold = bold
            r.font.name = "Calibri"
    return p


def main():
    doc = Document()

    # Tight margins so we actually fit on one page
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.1

    # --- Title block ---
    title = doc.add_paragraph()
    tight(title, space_after=0)
    r = title.add_run("DakCUPAC Bourbon & Whiskey Tasting Fundraiser")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    tight(sub, space_after=4)
    r = sub.add_run("CEO briefing  |  Target: Thursday evening, late summer / early fall 2026  |  Bismarck, ND")
    r.font.size = Pt(9.5)
    r.italic = True
    r.font.color.rgb = BLUE

    # --- The opportunity (one line) ---
    p = doc.add_paragraph()
    tight(p, space_after=4)
    r = p.add_run(
        "A private, 90-minute guided bourbon tasting for 40 CU executives, board members, and CU-affiliated "
        "PAC donors — led by Scott “Fuzzy” Meske. $250/ticket, six pours, engraved Glencairn glass gifted at "
        "check-in, Chick-fil-A catering. Net $8,500+ to PAC on $10,000 gross."
    )
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    # --- Financial snapshot table ---
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    widths = [Inches(1.85), Inches(1.85), Inches(1.85), Inches(1.85)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            set_cell_margins(cell, top=60, bottom=60, left=90, right=90)

    headers = ["Gross revenue", "Glassware", "Catering", "Net to PAC"]
    values = ["$10,000", "$763.20", "$480 – $700", "$8,537 – $8,757"]
    subs = ["40 × $250", "48 engraved Glencairn @ $15.90", "Chick-fil-A, spicy + non-spicy", "bottles, venue, Scott’s fee carried separately"]

    for i, (h, v, s) in enumerate(zip(headers, values, subs)):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "1F2A44")
        p = cell.paragraphs[0]
        tight(p, space_after=0)
        r = p.add_run(h.upper())
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        cell = table.rows[1].cells[i]
        shade_cell(cell, "F3F5F9")
        p = cell.paragraphs[0]
        tight(p, space_after=0)
        r = p.add_run(v)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = NAVY
        p2 = cell.add_paragraph()
        tight(p2, space_after=0)
        r = p2.add_run(s)
        r.font.size = Pt(8)
        r.font.color.rgb = LIGHT
        r.italic = True

    # --- Why it works ---
    add_heading(doc, "Why it works", space_before=8)
    add_bullet(doc, [
        ("Right audience, right price. ", True, False, None),
        ("Credit union execs and board members are precisely the donor pool with room to give at the PAC contribution ceiling. $250 reads as a tasting ticket, not a fundraiser ask.", False, False, None),
    ])
    add_bullet(doc, [
        ("Sponsorships eat the biggest cost line. ", True, False, None),
        ("Six CU bottle sponsors cover the whiskey, so PAC keeps nearly all of the gate.", False, False, None),
    ])
    add_bullet(doc, [
        ("The glass is the marketing. ", True, False, None),
        ("Every attendee leaves with an engraved Glencairn that lives in their bar. DakCUPAC shows up every time they pour for the next decade.", False, False, None),
    ])

    # --- Sponsorship ---
    add_heading(doc, "Sponsorship — six CU bottle sponsors (~$500/bottle)")
    add_bullet(doc, [
        ("Each sponsor: ", True, False, None),
        ("tent card next to their pour, verbal callout from Scott, logo on menu + check-in signage, two complimentary tickets (worth $500).", False, False, None),
    ])
    add_bullet(doc, [
        ("Optional Featured ND Pour sponsor ", True, False, None),
        ("underwrites a Heart River Spirits Painted Canyon (Bismarck craft distiller) host pour.", False, False, None),
    ])

    # --- Bottle lineup ---
    add_heading(doc, "Bottle lineup — each buyable today around $500")
    add_bullet(doc, [
        ("", False, True, "https://canawineco.com/products/2026-michters-10-year-old-single-barrel-bourbon-whiskey-750ml"),
        ("Michter’s 10 Yr Bourbon (2026)", False, True, "https://canawineco.com/products/2026-michters-10-year-old-single-barrel-bourbon-whiskey-750ml"),
        (" – $550  |  ", False, False, None),
        ("Old Forester Birthday 2025", False, True, "https://huntsmanheritage.com/products/old-forester-birthday-bourbon-2025"),
        (" – $700  |  ", False, False, None),
        ("Macallan 18 Sherry Oak", False, True, "https://flaviar.com/products/the-macallan-18-year-old-sherry-oak-2025-release-single-malt-scotch-whisky-750"),
        (" – $412", False, False, None),
    ])
    # Remove the duplicate first anchor that the list form required — simpler to just re-run cleanly:
    # (the bullet helper doesn't de-dupe, so build a second bullet instead)
    add_bullet(doc, [
        ("", False, False, None),
        ("Redbreast 21 Yr Irish", False, True, "https://www.totalwine.com/spirits/irish-whiskey/redbreast-21-yr-irish-whiskey/p/137936750"),
        (" – $400–500  |  ", False, False, None),
        ("Willett Family Estate 8 Yr", False, True, "https://www.blackwellswines.com/products/willett-family-estate-8-year-bourbon-whiskey"),
        (" – $300–550  |  ", False, False, None),
        ("Parker’s Heritage", False, True, "https://thebourbonconcierge.com/collections/parkers-heritage"),
        (" – $500+", False, False, None),
    ])

    # --- Glassware + catering detail ---
    add_heading(doc, "Glassware & catering")
    add_bullet(doc, [
        ("48 engraved Glencairn glasses at ", False, False, None),
        ("Sip n Savor", False, True, "https://sipnsavor.co/products/copy-of-bulk-glencairn-whisky-glasses"),
        (" — $15.90 each at the 48+ tier = $763.20. Spares cover breakage + late RSVPs.", False, False, None),
    ])
    add_bullet(doc, [
        ("Chick-fil-A catering at ", False, False, None),
        ("Kirkwood Mall Bismarck", False, True, "https://www.chick-fil-a.com/locations/nd/kirkwood-mall-nd"),
        (" (701-751-0793, $175 delivery min): 2× ", False, False, None),
        ("Large Nugget Tray", False, True, "https://www.chick-fil-a.com/catering/catering-trays/large-hot-chick-fil-a-nuggets-tray"),
        (", classic + spicy sandwich trays, mac & cheese, fruit. Budget $480–$700.", False, False, None),
    ])

    # --- Timeline ---
    add_heading(doc, "Timeline")
    add_bullet(doc, [
        ("T-12: ", True, False, None),
        ("date + venue locked, Chick-fil-A quote secured.  ", False, False, None),
        ("T-10: ", True, False, None),
        ("sponsor outreach opens, glass artwork submitted.  ", False, False, None),
        ("T-8: ", True, False, None),
        ("sponsorships in.  ", False, False, None),
        ("T-3: ", True, False, None),
        ("glass delivery.  ", False, False, None),
        ("T-2: ", True, False, None),
        ("final headcount. Day-of: doors 5:30, tasting 6:00.", False, False, None),
    ])

    # --- The decision ---
    add_heading(doc, "The decision", color=NAVY)
    add_bullet(doc, [
        ("Approve $250/40 anchor and ~$1,500 hard cost (glasses + catering) against $10K gross.", False, False, None),
    ])
    add_bullet(doc, [
        ("Green-light chair to open sponsor conversations with six CUs and book date with Scott.", False, False, None),
    ])
    add_bullet(doc, [
        ("Confirm venue lane: Elks Lodge #1199 (default) or a host residence if one volunteers.", False, False, None),
    ])

    out_path = "/home/user/dakcupac-2026/events/bourbon-tasting-ceo-brief.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
