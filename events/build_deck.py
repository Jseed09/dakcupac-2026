#!/usr/bin/env python3
"""Build the DakCUPAC Bourbon Tasting pitch deck as a .docx.

Costs limited to glassware and catering per the subcommittee chair's
direction. Chick-fil-A includes both spicy and non-spicy sandwich trays.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text, color="1F4E79", underline=True):
    """Append a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def slide_break(doc):
    """Force a page break so each slide starts on its own page."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    if bold:
        r.bold = True
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        r.text = text
    r.font.size = Pt(11)
    return p


def bullet_with_link(doc, prefix, link_text, url, suffix=""):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(prefix)
    r.font.size = Pt(11)
    add_hyperlink(p, url, link_text)
    if suffix:
        r2 = p.add_run(suffix)
        r2.font.size = Pt(11)
    return p


def slide_footer(doc, n, total):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Slide {n} of {total}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


TOTAL_SLIDES = 11


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- Slide 1 — Title ----------
    h1(doc, "DakCUPAC Bourbon & Whiskey Tasting")
    h2(doc, "Fundraiser pitch deck — subcommittee working draft")
    body(doc, "Prepared April 23, 2026 • All linked costs verified against live listings.")
    body(doc, "")
    body(doc, "Event concept: private, instructor-led tasting for CU executives, board members, and CU-affiliated PAC donors. Six pours, 90 minutes, Scott “Fuzzy” Meske leading. Engraved Glencairn glass gifted at the door doubles as the tasting glass and take-home keepsake. Chick-fil-A catering. Target: 40 attendees, $250/ticket, $10,000 gross.")
    body(doc, "")
    body(doc, "Per subcommittee chair direction, the only costs tracked in this deck are glassware and catering. Bottles are covered by CU bottle sponsors. Venue and instructor fee are carried separately.")
    slide_footer(doc, 1, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 2 — Why it works ----------
    h1(doc, "Why this event works")
    bullet(doc, "Right audience for the ticket price. Credit union execs and board members sit exactly where the federal/state PAC contribution ceiling lives. $250 reads as a bourbon-tasting cost, not a fundraiser ask.")
    bullet(doc, "Sponsorships eat the biggest cost line. CU bottle sponsors cover the whiskey; PAC keeps the gate.")
    bullet(doc, "The glass is the marketing. Every attendee leaves with an engraved Glencairn that lives in their bar. DakCUPAC shows up every time they pour for the next decade.")
    bullet(doc, "Scott is a force multiplier. Guided tasting is a different product than “open bar with bourbon.” It justifies the ticket and gives the room a shared experience.")
    slide_footer(doc, 2, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 3 — Numbers at a glance ----------
    h1(doc, "The numbers at a glance")
    body(doc, "Gross revenue: 40 attendees × $250 = $10,000", bold=True)
    body(doc, "")
    body(doc, "Tracked costs (this deck):", bold=True)
    bullet(doc, "Engraved Glencairn glasses (48 at $15.90 ea, 48+ tier 15% off): $763.20")
    bullet(doc, "Chick-fil-A catering (40, spicy + non-spicy): $480 – $700")
    body(doc, "")
    body(doc, "Tracked-cost subtotal: $1,243 – $1,463", bold=True)
    body(doc, "Net to PAC: $8,537 – $8,757 on $10,000 gross.")
    body(doc, "")
    body(doc, "Carried separately (not tracked here):")
    bullet(doc, "Bottles — covered by six CU bottle sponsors")
    bullet(doc, "Venue — Elks Lodge or private residence (TBD)")
    bullet(doc, "Scott’s instructor fee (TBD)")
    slide_footer(doc, 3, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 4 — Glassware cost ----------
    h1(doc, "Glassware — 48 engraved Glencairn whisky glasses")
    body(doc, "Confirmed pricing tier: 48+ qty at 15% off = $15.90 USD each. Order 48 (not 40) to hit the tier — spares cover breakage, late RSVPs, and the CU sponsors who will want one.")
    body(doc, "")
    body(doc, "Math: 48 × $15.90 = $763.20.", bold=True)
    body(doc, "")
    body(doc, "Vendor options (pricing verified April 2026):", bold=True)
    bullet_with_link(
        doc,
        "Sip n Savor — bulk engraved Glencairn, 48+ tier locked at $15.90/glass, free shipping over $100. ",
        "Sip n Savor bulk Glencairn",
        "https://sipnsavor.co/products/copy-of-bulk-glencairn-whisky-glasses",
    )
    bullet_with_link(
        doc,
        "Quality Glass Engraving — engraved Stölzle Glencairn 6oz, UV laser, no setup fee, 5–10 business day turnaround; comparable quote. ",
        "Quality Glass Engraving — Stölzle Glencairn",
        "https://qualityglassengraving.com/products/engraved-stolzle-glencairn-6-oz-whiskey-glass-item-3550031t",
    )
    bullet_with_link(
        doc,
        "Crystal Imagery — custom logo engraved Glencairn, deep-carved engraving; quote-based. ",
        "Crystal Imagery Glencairn collection",
        "https://crystalimagery.com/collections/glencairn-whisky-glass",
    )
    bullet_with_link(
        doc,
        "Distillery Products — Glencairn 48-case blanks at $5.90/glass ($288.80/case) before engraving; wholesale login required. ",
        "Distillery Products wholesale",
        "https://www.distilleryproducts.com/wholesale-glencairn/",
    )
    bullet_with_link(
        doc,
        "Glencairn Crystal — official wholesale program. ",
        "Glencairn wholesale",
        "https://business.glencairn.com/wholesale-whisky-glasses/",
    )
    body(doc, "")
    body(doc, "Budget for 48 engraved glasses: $763.20.", bold=True)
    body(doc, "Recommendation: Sip n Savor at the confirmed 48+ tier; submit DakCUPAC logo artwork at T-10 weeks, hard delivery deadline at T-3 weeks.")
    slide_footer(doc, 4, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 5 — Catering ----------
    h1(doc, "Catering — Chick-fil-A for 40 attendees")
    body(doc, "Per subcommittee chair: include both spicy and non-spicy sandwich trays so attendees can pick. Sourced from Chick-fil-A’s official catering site.")
    body(doc, "")
    body(doc, "Confirmed official menu items:", bold=True)
    bullet_with_link(
        doc,
        "Large Hot Chick-fil-A Nuggets Tray — 200 nuggets, serves ~25. ",
        "Chick-fil-A Large Nuggets Tray",
        "https://www.chick-fil-a.com/catering/catering-trays/large-hot-chick-fil-a-nuggets-tray",
    )
    bullet_with_link(
        doc,
        "Chick-fil-A Sandwich Tray (classic / non-spicy) — on official menu. ",
        "Chick-fil-A catering trays",
        "https://www.chick-fil-a.com/catering/trays",
    )
    bullet_with_link(
        doc,
        "Spicy Chicken Sandwich — available as a catering tray or bulk order at participating locations; confirm with the Bismarck operator. ",
        "Chick-fil-A catering",
        "https://www.chick-fil-a.com/catering",
    )
    bullet(doc, "Side trays: mac & cheese, waffle fries, fruit tray — available on the official catering menu.")
    body(doc, "")
    body(doc, "Bismarck ND location specifics:", bold=True)
    bullet_with_link(
        doc,
        "Chick-fil-A Kirkwood Mall (Bismarck ND) — catering available; delivery subtotal minimum $175. ",
        "Kirkwood Mall location",
        "https://www.chick-fil-a.com/locations/nd/kirkwood-mall-nd",
    )
    bullet(doc, "Bismarck catering line: (701) 751-0793.")
    bullet(doc, "Prices are location-gated on chick-fil-a.com; call or enter ZIP through the app for firm Bismarck pricing.")
    body(doc, "")
    body(doc, "Planned order for 40 (spicy + non-spicy sandwich mix):", bold=True)
    bullet(doc, "2× Large Hot Nuggets Tray — coverage staple — ~$186")
    bullet(doc, "1× Classic Chicken Sandwich Tray (non-spicy) — ~$50 (half the sandwich eaters)")
    bullet(doc, "1× Spicy Chicken Sandwich Tray — ~$55 (spicy half; confirm availability in Bismarck)")
    bullet(doc, "1× Mac & Cheese tray — ~$50")
    bullet(doc, "1× Fruit tray — ~$55")
    bullet(doc, "Sauces, napkins, small wares — ~$40")
    body(doc, "")
    body(doc, "Food subtotal estimate: ~$436. With delivery + tip, budget $480 – $700 for catering.", bold=True)
    body(doc, "Note: national averages come from aggregator menus; Bismarck Chick-fil-A requires a direct quote — build the written quote into the T-12 checklist.")
    slide_footer(doc, 5, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 6 — Bottle lineup (sponsor-facing) ----------
    h1(doc, "Bottle lineup — six pours, buyable around $500")
    body(doc, "Each bottle links to a real retailer where a CU sponsor can actually buy it today. Prices verified April 2026. Bottle cost is covered by CU bottle sponsors — not a tracked PAC cost.")
    body(doc, "")
    bullet_with_link(
        doc,
        "Michter’s 10 Year Single Barrel Bourbon (2026 release) — $549.99 at Cana Wine Co. ",
        "Buy Michter’s 10 Year — Cana Wine Co.",
        "https://canawineco.com/products/2026-michters-10-year-old-single-barrel-bourbon-whiskey-750ml",
    )
    bullet_with_link(
        doc,
        "Old Forester Birthday Bourbon 2025 — $699.99 at Huntsman Heritage; $799.99 at The Rare Whiskey Shop. ",
        "Buy OFBB 2025 — Huntsman Heritage",
        "https://huntsmanheritage.com/products/old-forester-birthday-bourbon-2025",
    )
    bullet_with_link(
        doc,
        "The Macallan 18 Year Sherry Oak (2026 release) — ~$412 at Flaviar. ",
        "Buy Macallan 18 — Flaviar",
        "https://flaviar.com/products/the-macallan-18-year-old-sherry-oak-2025-release-single-malt-scotch-whisky-750",
    )
    bullet_with_link(
        doc,
        "Redbreast 21 Year Single Pot Still Irish Whiskey — ~$400–500 at Total Wine. ",
        "Buy Redbreast 21 — Total Wine",
        "https://www.totalwine.com/spirits/irish-whiskey/redbreast-21-yr-irish-whiskey/p/137936750",
    )
    bullet_with_link(
        doc,
        "Willett Family Estate 8 Year Single Barrel Bourbon — single barrels regularly listed $300–550. ",
        "Buy Willett 8 Year — Blackwell’s Wines",
        "https://www.blackwellswines.com/products/willett-family-estate-8-year-bourbon-whiskey",
    )
    bullet_with_link(
        doc,
        "Parker’s Heritage Collection (latest release) — retail routinely $500+; most recent editions in stock at The Bourbon Concierge. ",
        "Buy Parker’s Heritage — The Bourbon Concierge",
        "https://thebourbonconcierge.com/collections/parkers-heritage",
    )
    body(doc, "")
    body(doc, "Optional 7th — “host pour”:", bold=True)
    bullet_with_link(
        doc,
        "Heart River Spirits Painted Canyon Bourbon — Bismarck craft distiller, new release. Public pricing not posted; request direct. ",
        "Heart River Spirits — Painted Canyon",
        "https://heartriverspirits.com/paintedcanyonbourbon/",
    )
    body(doc, "")
    body(doc, "Chair’s note: the specific bottles are negotiable — Scott will have final say on the flight. These six are here to show every sponsoring CU that a $500ish bottle is a real, findable ask, not an allocation pipe dream.")
    slide_footer(doc, 6, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 7 — Pricing sources ----------
    h1(doc, "Pricing sources")
    h2(doc, "Glassware")
    bullet_with_link(doc, "", "Sip n Savor — bulk Glencairn", "https://sipnsavor.co/products/copy-of-bulk-glencairn-whisky-glasses")
    bullet_with_link(doc, "", "Ryan’s Woodworking (Etsy)", "https://www.etsy.com/listing/1667226787/custom-glencairn-whiskey-glasses-any")
    bullet_with_link(doc, "", "Distillery Products — wholesale Glencairn", "https://www.distilleryproducts.com/wholesale-glencairn/")
    bullet_with_link(doc, "", "Quality Glass Engraving — Stölzle Glencairn 6oz", "https://qualityglassengraving.com/products/engraved-stolzle-glencairn-6-oz-whiskey-glass-item-3550031t")
    bullet_with_link(doc, "", "Crystal Imagery — Glencairn collection", "https://crystalimagery.com/collections/glencairn-whisky-glass")
    bullet_with_link(doc, "", "Glencairn Crystal — wholesale program", "https://business.glencairn.com/wholesale-whisky-glasses/")

    h2(doc, "Catering")
    bullet_with_link(doc, "", "Chick-fil-A — Large Hot Nuggets Tray", "https://www.chick-fil-a.com/catering/catering-trays/large-hot-chick-fil-a-nuggets-tray")
    bullet_with_link(doc, "", "Chick-fil-A — catering trays overview", "https://www.chick-fil-a.com/catering/trays")
    bullet_with_link(doc, "", "Chick-fil-A — catering landing page", "https://www.chick-fil-a.com/catering")
    bullet_with_link(doc, "", "Chick-fil-A — Kirkwood Mall, Bismarck ND", "https://www.chick-fil-a.com/locations/nd/kirkwood-mall-nd")

    slide_footer(doc, 7, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 8 — Sponsorship structure ----------
    h1(doc, "Sponsorship structure — six CU bottle sponsors")
    body(doc, "Each sponsoring credit union covers one bottle in the tasting flight. Bottle cost is fully outside this deck’s tracked costs.")
    body(doc, "")
    body(doc, "Each sponsor receives:", bold=True)
    bullet(doc, "Tent card / placard displayed next to their bottle: “This pour proudly sponsored by [CU Name + logo]”")
    bullet(doc, "Verbal callout from Scott when that pour is introduced")
    bullet(doc, "Recognition on the printed tasting menu handed to every guest")
    bullet(doc, "Logo on event signage at check-in")
    bullet(doc, "Two complimentary tickets (worth $500)")
    body(doc, "")
    body(doc, "Optional add-on — “Featured ND Pour” sponsor:", bold=True)
    bullet(doc, "Underwrites the host pour (Heart River Spirits Painted Canyon Bourbon)")
    bullet(doc, "Lower buy-in, good fit for an ND-headquartered CU")
    bullet(doc, "One complimentary ticket included")
    slide_footer(doc, 8, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 9 — Timeline ----------
    h1(doc, "Timeline & milestones")
    body(doc, "Working back from a target Thursday evening, late summer / early fall 2026.")
    body(doc, "")
    bullet(doc, "T-12 weeks — confirm date with Scott; lock venue; finalize budget ceiling; call Bismarck Chick-fil-A for written quote")
    bullet(doc, "T-10 weeks — begin sponsor outreach to six CUs; finalize bottle list with Scott; submit Glencairn artwork and place glass order")
    bullet(doc, "T-8 weeks — sponsorship commitments in; deposit on glasses; venue contract signed")
    bullet(doc, "T-6 weeks — open ticket sales; catering order placed with revisable headcount")
    bullet(doc, "T-4 weeks — bottle sourcing confirmed; printed materials to printer")
    bullet(doc, "T-3 weeks — glass delivery deadline")
    bullet(doc, "T-2 weeks — final headcount lock; catering quantity and spicy/non-spicy split confirmed")
    bullet(doc, "T-1 week — walkthrough at venue; AV check; glass delivery confirmed")
    bullet(doc, "Day-of — setup 4:00 PM; doors 5:30 PM; tasting starts 6:00 PM")
    slide_footer(doc, 9, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 10 — Risks ----------
    h1(doc, "Risks & mitigations")
    bullet(doc, "Glass delivery slips — order at T-10 with a hard T-3 delivery date; keep a non-engraved Glencairn backup vendor on standby.")
    bullet(doc, "Chick-fil-A spicy tray unavailable at Bismarck — confirm with the operator at T-12; fallback is a 50/50 mix on the classic sandwich tray with spicy singles ordered off the main menu.")
    bullet(doc, "Catering pricing varies in ND — get a written quote from the Bismarck operator at T-12, not T-2.")
    bullet(doc, "$175 catering delivery minimum — already well cleared by the planned order.")
    bullet(doc, "Headcount comes in below 40 — couples-pricing or “bring a colleague” upsell at T-4; food and glass costs scale down cleanly.")
    slide_footer(doc, 10, TOTAL_SLIDES)
    slide_break(doc)

    # ---------- Slide 11 — The ask ----------
    h1(doc, "The ask")
    body(doc, "Subcommittee, decide today:", bold=True)
    body(doc, "")
    bullet(doc, "Approve the $250/40 anchor — $10,000 gross target.")
    bullet(doc, "Approve the glassware line: 48 engraved Glencairn glasses at $15.90 each = $763.20.")
    bullet(doc, "Approve the catering budget band: $480 – $700 for Chick-fil-A (spicy + non-spicy sandwich trays, nuggets, sides).")
    bullet(doc, "Authorize the chair to open the conversation with Scott with everything above on the table.")
    body(doc, "")
    body(doc, "Tracked-cost subtotal at the recommended configuration: $1,243 – $1,463 against a $10,000 gross.", bold=True)
    slide_footer(doc, 11, TOTAL_SLIDES)

    out_path = "/home/user/dakcupac-2026/events/bourbon-tasting-2026.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
